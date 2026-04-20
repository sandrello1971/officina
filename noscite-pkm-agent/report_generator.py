"""
report_generator.py — Genera un report ragionato su un insieme di documenti selezionati.
Legge gli estratti dai file .md in _metadata/ e chiama Claude per l'analisi.
"""

import logging
from datetime import datetime
from pathlib import Path

import anthropic

from config import Config
from api_utils import call_with_retry

log = logging.getLogger(__name__)
client = anthropic.Anthropic()

SYSTEM_PROMPT = """Sei un analista esperto di knowledge management.
Ricevi una lista di documenti con i loro estratti e metadati.
Produci un report ragionato in italiano con ESATTAMENTE queste sezioni nell'ordine indicato.
Usa titoli markdown ## per ogni sezione principale e ### per le sottosezioni."""


def generate(stems: list[str], custom_prompt: str = "", focus_items: list[str] = []) -> dict | None:
    """
    Genera un report ragionato per i documenti con gli stem indicati.
    Restituisce dict con report_stem, tokens, o None in caso di errore.
    """
    if not stems:
        return None

    # Leggi i metadati e gli estratti di ciascun documento
    docs_content = []
    for stem in stems:
        md_path = Config.METADATA_DIR / f"{stem}.md"
        if not md_path.exists():
            log.warning(f"  ⚠ File non trovato: {stem}.md")
            continue
        text = md_path.read_text(encoding="utf-8", errors="replace")
        info = _parse_doc_info(stem, text)
        docs_content.append(info)

    if not docs_content:
        log.error("  ✗ Nessun documento trovato per il report")
        return None

    log.info(f"  📊 Generazione report per {len(docs_content)} documenti...")

    # Costruisci il blocco documenti
    docs_block = ""
    for i, doc in enumerate(docs_content, 1):
        docs_block += f"\n## Documento {i}: {doc['title']}\n"
        docs_block += f"**Tipo:** {doc['tipo']} | **Lingua:** {doc['lingua']} | **Data:** {doc['data']}\n"
        docs_block += f"**Sommario:** {doc['sommario']}\n\n"
        if doc['estratto']:
            docs_block += f"**Estratto:**\n{doc['estratto'][:2000]}\n"
        if doc['tags']:
            docs_block += f"\n**Tag:** {', '.join(doc['tags'][:10])}\n"
        docs_block += "\n---\n"

    # Sezione focus opzionale
    focus_block = ""
    if focus_items:
        focus_labels = {
            "persone": "Persone e figure chiave",
            "organizzazioni": "Organizzazioni e aziende",
            "territori": "Territori e luoghi",
            "tecnologie": "Tecnologie e strumenti",
            "normative": "Leggi e normative",
            "prodotti": "Prodotti e dispositivi",
            "date": "Date e timeline",
            "dati": "Dati e statistiche",
        }
        focus_str = ", ".join(focus_labels.get(f, f) for f in focus_items)
        focus_block = f"\nFocalizza l'analisi in modo approfondito su: **{focus_str}**.\n"

    custom_block = f"\nRichieste specifiche dell'utente: {custom_prompt}\n" if custom_prompt else ""

    msg = f"""Analizza i seguenti {len(docs_content)} documenti e produci un report strutturato.
{focus_block}{custom_block}

{docs_block}

Produci il report con ESATTAMENTE queste sezioni in questo ordine:

## 1. Documenti analizzati
Elenco puntato con: titolo, tipo, data, breve descrizione (1 frase) per ciascun documento.

## 2. Elementi comuni
Temi, concetti, tecnologie, entità o argomenti che compaiono in più documenti. Raggruppa per categoria e spiega le connessioni.{(' Includi sezioni specifiche per: ' + ', '.join(focus_items) + '.') if focus_items else ''}

## 3. Elementi discordanti
Differenze, contraddizioni, approcci opposti o elementi presenti in alcuni documenti ma assenti in altri. Spiega la rilevanza delle divergenze.

## 4. Analisi trasversale
Connessioni non ovvie, pattern emergenti, lacune informative, domande aperte che il corpus solleva.

## 5. Conclusioni e spunti
Sintesi del valore complessivo del corpus e 3-5 spunti concreti per approfondimenti futuri."""

    try:
        r = call_with_retry(
            client.messages.create,
            model=Config.CLAUDE_MODEL,
            max_tokens=6000,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": msg}],
        )
        report_text = r.content[0].text.strip()
        tokens_in  = r.usage.input_tokens
        tokens_out = r.usage.output_tokens
        log.info(f"  ✓ Report generato ({tokens_in} in / {tokens_out} out)")
    except Exception as e:
        log.error(f"  ✗ Errore generazione report: {e}")
        return None

    # Salva il report come .md
    now = datetime.now()
    stamp = now.strftime("%Y%m%d_%H%M%S")
    report_stem = f"_Report_{stamp}"
    report_path = Config.METADATA_DIR / f"{report_stem}.md"

    today = now.strftime("%Y-%m-%d")
    now_str = now.strftime("%d/%m/%Y %H:%M")
    stems_links = " · ".join(f"[[{s}]]" for s in stems)
    titoli = ", ".join(d["title"] for d in docs_content[:3])
    if len(docs_content) > 3:
        titoli += f" e altri {len(docs_content) - 3}"

    # Genera il .docx in _archive prima di scrivere il .md (serve il nome file)
    docx_path = generate_docx(report_text, stems, docs_content, report_stem)
    docx_link = f'"[[{report_stem}.docx]]"' if docx_path else "N/D"

    focus_labels = {
        "persone": "Persone", "organizzazioni": "Organizzazioni",
        "territori": "Territori e luoghi", "tecnologie": "Tecnologie",
        "normative": "Leggi e normative", "prodotti": "Prodotti e dispositivi",
        "date": "Date e timeline", "dati": "Dati e statistiche",
    }
    focus_str = ", ".join(focus_labels.get(f, f) for f in focus_items) if focus_items else ""

    # Righe opzionali nel frontmatter
    focus_yaml = f'focus_analisi: "{focus_str}"\n' if focus_str else ""
    prompt_yaml = f'prompt_integrativo: "{custom_prompt.replace(chr(34), chr(39))}"\n' if custom_prompt else ""

    # Righe opzionali nell'header del documento
    focus_header = f"> Focus analisi: {focus_str}\n" if focus_str else ""
    prompt_header = f"> Prompt integrativo: _{custom_prompt}_\n" if custom_prompt else ""

    frontmatter = f"""---
title: "Report — {titoli}"
tipo_documento: report_analisi
data_catalogazione: {today}
data_documento: {today}
documenti_analizzati: {len(docs_content)}
{focus_yaml}{prompt_yaml}sommario: "Report ragionato su {len(docs_content)} documenti: {titoli}."
file_originale: {docx_link}
lingua: it
tags:
  - report_analisi
  - pkm_agent
created: {now_str}
---

> **Report generato automaticamente da PKM Agent**
> Documenti analizzati: {stems_links}
> Data: {now_str} | Token: {tokens_in} in / {tokens_out} out
> File Word: {docx_link}
{focus_header}{prompt_header}
---

"""

    report_path.write_text(frontmatter + report_text, encoding="utf-8")
    log.info(f"  ✓ Report salvato: {report_stem}.md")

    return {
        "report_stem": report_stem,
        "tokens_input": tokens_in,
        "tokens_output": tokens_out,
        "doc_count": len(docs_content),
    }


def _parse_doc_info(stem: str, text: str) -> dict:
    """Estrae titolo, sommario, estratto e metadati dal testo del file .md."""
    info = {
        "stem": stem,
        "title": stem,
        "tipo": "N/D",
        "lingua": "N/D",
        "data": "N/D",
        "sommario": "",
        "estratto": "",
        "tags": [],
    }

    # Frontmatter
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            fm = text[3:end]
            for line in fm.splitlines():
                line = line.strip()
                if not line or ":" not in line:
                    continue
                key, _, val = line.partition(":")
                key = key.strip()
                val = val.strip().strip('"').strip("'")
                if key == "title":        info["title"] = val
                elif key == "tipo_documento": info["tipo"] = val
                elif key == "lingua":     info["lingua"] = val
                elif key == "data_documento": info["data"] = val
                elif key == "sommario":   info["sommario"] = val

            # Tags
            in_tags = False
            for line in fm.splitlines():
                stripped = line.strip()
                if stripped == "tags:":
                    in_tags = True
                    continue
                if in_tags:
                    if stripped.startswith("- "):
                        info["tags"].append(stripped[2:].strip().strip('"'))
                    elif stripped and not stripped.startswith("-"):
                        break

    # Estratto dal corpo
    marker = "## 📄 Estratto documento"
    idx = text.find(marker)
    if idx != -1:
        start = text.find("\n", idx) + 1
        next_h2 = text.find("\n## ", start)
        next_sep = text.find("\n---", start)
        candidates = [i for i in [next_h2, next_sep] if i != -1]
        end = min(candidates) if candidates else len(text)
        info["estratto"] = text[start:end].strip()

    return info


def generate_docx(report_text: str, stems: list[str], doc_infos: list[dict],
                  report_stem: str) -> Path | None:
    """
    Genera un file .docx professionale usando node/docx-js.
    Include sommario automatico e formattazione LCARS-inspired.
    """
    import subprocess
    import json
    import shutil
    import os as _os

    # Cerca node in percorsi comuni su macOS
    node_path = shutil.which("node")
    if not node_path:
        for candidate in ["/opt/homebrew/bin/node", "/usr/local/bin/node", "/usr/bin/node"]:
            if _os.path.exists(candidate):
                node_path = candidate
                break
    if not node_path:
        log.warning("  ⚠ node.js non disponibile, uso fallback python-docx")
        return _generate_docx_fallback(report_text, stems, doc_infos, report_stem,
                                       Config.ARCHIVE_DIR / f"{report_stem}.docx",
                                       datetime.now().strftime("%d/%m/%Y %H:%M"))

    log.info(f"  📄 Usando node: {node_path}")

    now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    docx_path = Config.ARCHIVE_DIR / f"{report_stem}.docx"

    # Prepara i dati da passare allo script node
    data = {
        "report_text": report_text,
        "doc_infos": doc_infos,
        "stems": stems,
        "now_str": now_str,
        "output_path": str(docx_path),
        "report_stem": report_stem,
    }

    # Script node inline
    js = r"""
const fs = require('fs');
const path = require('path');
const {
  Document, Packer, Paragraph, TextRun, Table, TableRow, TableCell,
  HeadingLevel, AlignmentType, BorderStyle, WidthType, ShadingType,
  LevelFormat, PageBreak, Footer, PageNumber,
  PageNumberElement
} = require('docx');

const data = JSON.parse(fs.readFileSync('/tmp/report_data.json', 'utf8'));
const BLUE   = "1A5C9A";
const ORANGE = "CC6600";
const GRAY   = "F2F2F2";
const MUTED  = "666666";
const border = { style: BorderStyle.SINGLE, size: 1, color: "CCCCCC" };
const borders = { top: border, bottom: border, left: border, right: border };

function h1(text) {
  return new Paragraph({
    heading: HeadingLevel.HEADING_1,
    spacing: { before: 400, after: 120 },
    border: { bottom: { style: BorderStyle.SINGLE, size: 8, color: BLUE, space: 1 } },
    children: [new TextRun({ text, bold: true, size: 32, color: BLUE, font: "Arial" })]
  });
}
function h2(text) {
  return new Paragraph({
    heading: HeadingLevel.HEADING_2,
    spacing: { before: 280, after: 80 },
    children: [new TextRun({ text, bold: true, size: 26, color: "333333", font: "Arial" })]
  });
}
function h3(text) {
  return new Paragraph({
    heading: HeadingLevel.HEADING_3,
    spacing: { before: 180, after: 60 },
    children: [new TextRun({ text, bold: true, size: 22, color: "555555", font: "Arial" })]
  });
}
function para(text, opts) {
  opts = opts || {};
  const parts = text.split(/\*\*(.*?)\*\*/g);
  const runs = parts.map((p, i) => new TextRun({
    text: p, bold: i % 2 === 1, size: 22, font: "Arial", ...opts
  }));
  return new Paragraph({ spacing: { before: 60, after: 60 }, children: runs });
}
function bullet(text) {
  return new Paragraph({
    numbering: { reference: "bullets", level: 0 },
    spacing: { before: 40, after: 40 },
    children: [new TextRun({ text, size: 22, font: "Arial" })]
  });
}
function space() {
  return new Paragraph({ spacing: { before: 80, after: 80 }, children: [new TextRun("")] });
}
function divider() {
  return new Paragraph({
    spacing: { before: 120, after: 120 },
    border: { bottom: { style: BorderStyle.SINGLE, size: 4, color: BLUE, space: 1 } },
    children: [new TextRun("")]
  });
}

// Converti markdown in elementi docx
function mdToElements(text) {
  const elements = [];
  const lines = text.split('\n');
  for (const line of lines) {
    if (line.startsWith('## ')) elements.push(h2(line.slice(3).trim()));
    else if (line.startsWith('### ')) elements.push(h3(line.slice(4).trim()));
    else if (line.startsWith('#### ')) elements.push(h3(line.slice(5).trim()));
    else if (line.startsWith('- ') || line.startsWith('* ')) elements.push(bullet(line.slice(2).trim()));
    else if (/^\d+\. /.test(line)) elements.push(bullet(line.replace(/^\d+\. /, '').trim()));
    else if (line.trim() === '---') elements.push(divider());
    else if (line.trim() === '') elements.push(space());
    else elements.push(para(line));
  }
  return elements;
}

// Tabella documenti analizzati
function docsTable() {
  const headerRow = new TableRow({
    tableHeader: true,
    children: ["Documento", "Tipo", "Lingua", "Data"].map(label =>
      new TableCell({
        borders,
        shading: { fill: "1A5C9A", type: ShadingType.CLEAR },
        margins: { top: 80, bottom: 80, left: 120, right: 120 },
        children: [new Paragraph({ children: [new TextRun({ text: label, bold: true, size: 20, color: "FFFFFF", font: "Arial" })] })]
      })
    )
  });
  const dataRows = data.doc_infos.map(d =>
    new TableRow({
      children: [d.title, d.tipo, d.lingua, d.data].map((val, i) =>
        new TableCell({
          borders,
          shading: { fill: i === 0 ? GRAY : "FFFFFF", type: ShadingType.CLEAR },
          margins: { top: 60, bottom: 60, left: 120, right: 120 },
          children: [new Paragraph({ children: [new TextRun({ text: val || "N/D", size: 19, font: "Arial", bold: i === 0 })] })]
        })
      )
    })
  );
  return new Table({
    width: { size: 9026, type: WidthType.DXA },
    columnWidths: [4000, 1500, 1200, 1326],
    rows: [headerRow, ...dataRows]
  });
}

const doc = new Document({
  numbering: {
    config: [{
      reference: "bullets",
      levels: [{ level: 0, format: LevelFormat.BULLET, text: "•", alignment: AlignmentType.LEFT,
        style: { paragraph: { indent: { left: 720, hanging: 360 } } } }]
    }]
  },
  styles: {
    default: { document: { run: { font: "Arial", size: 22 } } },
    paragraphStyles: [
      { id: "Heading1", name: "Heading 1", basedOn: "Normal", next: "Normal", quickFormat: true,
        run: { size: 32, bold: true, font: "Arial", color: BLUE },
        paragraph: { spacing: { before: 400, after: 120 }, outlineLevel: 0 } },
      { id: "Heading2", name: "Heading 2", basedOn: "Normal", next: "Normal", quickFormat: true,
        run: { size: 26, bold: true, font: "Arial", color: "333333" },
        paragraph: { spacing: { before: 280, after: 80 }, outlineLevel: 1 } },
      { id: "Heading3", name: "Heading 3", basedOn: "Normal", next: "Normal", quickFormat: true,
        run: { size: 22, bold: true, font: "Arial", color: "555555" },
        paragraph: { spacing: { before: 180, after: 60 }, outlineLevel: 2 } },
    ]
  },
  sections: [{
    properties: {
      page: {
        size: { width: 11906, height: 16838 },
        margin: { top: 1134, right: 1134, bottom: 1134, left: 1134 }
      }
    },
    footers: {
      default: new Footer({
        children: [new Paragraph({
          alignment: AlignmentType.CENTER,
          children: [
            new TextRun({ text: "PKM Agent — Report di Analisi  |  ", size: 17, color: MUTED, font: "Arial" }),
            new TextRun({ children: [new PageNumberElement()], size: 17, color: MUTED, font: "Arial" })
          ]
        })]
      })
    },
    children: [
      // Titolo
      new Paragraph({
        alignment: AlignmentType.CENTER,
        spacing: { before: 480, after: 160 },
        children: [new TextRun({ text: "Report di Analisi", bold: true, size: 52, color: BLUE, font: "Arial" })]
      }),
      new Paragraph({
        alignment: AlignmentType.CENTER,
        spacing: { before: 0, after: 80 },
        children: [new TextRun({ text: `${data.doc_infos.length} documenti analizzati  ·  ${data.now_str}`, size: 22, color: MUTED, font: "Arial" })]
      }),
      divider(),
      space(),

      // Indice manuale — estratto dai titoli ## del report
      new Paragraph({
        heading: HeadingLevel.HEADING_1,
        spacing: { before: 200, after: 120 },
        children: [new TextRun({ text: "Indice", bold: true, size: 32, color: BLUE, font: "Arial" })]
      }),
      ...(() => {
        const tocItems = [];
        const lines = data.report_text.split('\n');
        let sectionNum = 0;
        for (const line of lines) {
          if (line.startsWith('## ')) {
            sectionNum++;
            const title = line.slice(3).trim();
            tocItems.push(new Paragraph({
              spacing: { before: 40, after: 40 },
              children: [
                new TextRun({ text: `${sectionNum}.  ${title}`, size: 22, font: "Arial", color: BLUE }),
              ]
            }));
          } else if (line.startsWith('### ')) {
            const title = line.slice(4).trim();
            tocItems.push(new Paragraph({
              spacing: { before: 20, after: 20 },
              indent: { left: 480 },
              children: [
                new TextRun({ text: `— ${title}`, size: 20, font: "Arial", color: "666666" }),
              ]
            }));
          }
        }
        return tocItems;
      })(),
      new Paragraph({ children: [new PageBreak()] }),

      // Corpo report
      ...mdToElements(data.report_text),

      space(),
      divider(),
      new Paragraph({
        alignment: AlignmentType.CENTER,
        spacing: { before: 120, after: 0 },
        children: [new TextRun({ text: "Report generato automaticamente da PKM Agent con Claude AI", size: 18, color: MUTED, font: "Arial", italics: true })]
      }),
    ]
  }]
});

Packer.toBuffer(doc).then(buf => {
  fs.writeFileSync(data.output_path, buf);
  console.log('OK');
});
"""

    try:
        # Scrivi i dati JSON
        with open("/tmp/report_data.json", "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)

        # Scrivi lo script node
        with open("/tmp/gen_report.js", "w", encoding="utf-8") as f:
            f.write(js)

        import os as _os
        # Rileva NODE_PATH dinamicamente da npm
        npm_root = "/opt/homebrew/lib/node_modules"
        try:
            nr = subprocess.run(["npm", "root", "-g"], capture_output=True, text=True, timeout=5)
            if nr.returncode == 0 and nr.stdout.strip():
                npm_root = nr.stdout.strip()
        except Exception:
            pass

        env = {**_os.environ, "NODE_PATH": npm_root}
        result = subprocess.run(
            [node_path, "/tmp/gen_report.js"],
            capture_output=True, text=True, timeout=60,
            env=env
        )

        log.info(f"  📄 NODE_PATH: {npm_root}")
        log.info(f"  📄 node stdout: {result.stdout[:200]}")
        if result.stderr:
            log.warning(f"  📄 node stderr: {result.stderr[:300]}")

        if result.returncode == 0 and "OK" in result.stdout:
            log.info(f"  ✓ Report DOCX salvato: {report_stem}.docx")
            return docx_path
        else:
            log.error(f"  ✗ Errore node DOCX (rc={result.returncode}): {result.stderr[:500]}")
            log.info("  → Fallback a python-docx...")
            return _generate_docx_fallback(report_text, stems, doc_infos, report_stem, docx_path, now_str)
    except Exception as e:
        log.error(f"  ✗ Errore generazione DOCX: {e}")
        return None


def _generate_docx_fallback(report_text, stems, doc_infos, report_stem, docx_path, now_str):
    """Fallback con python-docx se node non disponibile."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)
        t = doc.add_heading("Report di Analisi PKM", level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(f"Generato il {now_str}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        for line in report_text.splitlines():
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph()
                p.add_run("• " + line[2:].strip())
                p.paragraph_format.left_indent = Pt(18)
            elif line.strip() == "---":
                doc.add_paragraph("─" * 50)
            elif line.strip():
                # gestione bold **testo**
                p = doc.add_paragraph()
                parts = line.split("**")
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.bold = (i % 2 == 1)
        doc.save(str(docx_path))
        log.info(f"  ✓ Report DOCX (fallback) salvato: {report_stem}.docx")
        return docx_path
    except Exception as e:
        log.error(f"  ✗ Fallback DOCX fallito: {e}")
        return None
